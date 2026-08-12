import os
import json
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from typing import Optional
from pypdf import PdfReader
from docx import Document

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API Error")

client = Groq(api_key=my_api_key)

model = "llama-3.3-70b-versatile"   # swapped mid-lecture due to the primary model's rate limit

# ============================================================
# PART 1 — Job Description Schema & Extraction
# ============================================================

class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: str = None
    educational_requirements: list[str]
    responsibilities: list[str]

job_schema = JobD.model_json_schema()

job_description_text = """## Job Summary

We are looking for a motivated AI Engineer with *1 year of experience* in Machine Learning and Artificial Intelligence. The ideal candidate should have hands-on experience in developing, training, and deploying basic AI/ML models. The candidate will work closely with senior engineers to build AI-powered applications and continuously learn new technologies.

## Key Responsibilities

* Develop and train machine learning and deep learning models under guidance.

* Perform data collection, cleaning, preprocessing, and feature engineering.

* Fine-tune pre-trained models such as BERT, ResNet, or GPT for specific tasks.

* Assist in developing NLP and Computer Vision applications.

* Support the development of RAG-based chatbots using LangChain and vector databases.

* Work with LLMs and AI models such as *GPT, Claude, Gemini, Llama, and DeepSeek* through APIs or open-source frameworks.

* Perform exploratory data analysis (EDA) and prepare datasets for training.

* Evaluate model performance and improve accuracy using appropriate techniques.

* Deploy AI models using basic Docker containers and cloud services.

* Collaborate with software developers to integrate AI models into applications.

* Use Git for version control and maintain project documentation.

* Stay updated with the latest AI/ML tools and technologies.

## Requirements

* *1 year of experience* in AI, Machine Learning, or Deep Learning.

* Good knowledge of Python and AI libraries such as TensorFlow, PyTorch, or Scikit-learn.

* Basic understanding of machine learning algorithms including classification, regression, clustering, and neural networks.

* Familiarity with NLP or Computer Vision concepts.

* Basic knowledge of Generative AI, prompt engineering, and LLM applications.

* Exposure to LangChain, LangGraph, MCP, or RAG-based applications is a plus.

* Familiarity with AI models such as *GPT, Claude, Gemini, Llama, and DeepSeek* is preferred.

* Familiarity with Git and Docker.

* Basic understanding of REST APIs and model deployment.

* Good analytical and problem-solving skills.

* Strong communication and teamwork abilities.

* Ability to learn quickly and work in a collaborative environment.

## Preferred Skills

* Experience integrating AI APIs from OpenAI, Anthropic (Claude), Google Gemini, or open-source Llama models.

* Knowledge of vector databases such as FAISS, ChromaDB, or Pinecone.

* Familiarity with cloud platforms like AWS, Azure, or Google Cloud is a plus.

* Personal or academic AI projects demonstrating practical experience.

Pay: ₹15,000.00 - ₹20,000.00 per month

Benefits:

Paid sick time
Work from home
Application Question(s):

Preferred only Married
Work Location: Remote

"""

job_system_prompt = f"""
You are an expert HR assistant. Analyze the job description and extract
structured information. Return only valid JSON matching this schema: {job_schema}
Do not return the schema itself. Do not return unnecessary fields
(property, title, type). Fill the schema with actual extracted information.
If minimum experience is not mentioned, set it to null.
Do not invent information.
"""

job_user_prompt = f"Analyze the following job description: {job_description_text}"

message_system_job = {"role": "system", "content": job_system_prompt}
message_user_job = {"role": "user", "content": job_user_prompt}

messages_job = [message_system_job, message_user_job]

response_format = {"type": "json_object"}

response_job = client.chat.completions.create(
    model=model,
    messages=messages_job,
    response_format=response_format
)

raw_job_json = response_job.choices[0].message.content
job_data = json.loads(raw_job_json)
job = JobD(**job_data)

# ============================================================
# PART 2 — Résumé Schema Design
# ============================================================

class Experience(BaseModel):
    company_name: Optional[str] = None
    role: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[str] = None
    skills_used: list[str] = []


class Resume(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    total_experience_years: Optional[str] = None
    skills: list[str] = []
    experiences: list[Experience] = []
    projects: list[str] = []
    certifications: list[str] = []

resume_schema = Resume.model_json_schema()

# ============================================================
# PART 3 — Reading PDF and DOCX Files
# ============================================================

def read_pdf(file_path):
    text = ""
    reader = PdfReader(file_path)
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text

def read_docx(file_path):
    text = ""
    document = Document(file_path)
    for paragraph in document.paragraphs:
        text += paragraph.text

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text

    return text

def read_resume(file_path):
    if str(file_path).lower().endswith(".pdf"):
        return read_pdf(file_path)
    elif str(file_path).lower().endswith(".docx"):
        return read_docx(file_path)
    else:
        return None

# ============================================================
# PART 4 — Parsing Résumés
# ============================================================

def parse_resume(resume_text):
    resume_system_prompt = f"""
    You are an expert résumé parser. Extract information from the résumé.
    Different résumés use different headings — treat "Experience",
    "Professional Experience", "Work History", "Employment", and
    "Internship" as referring to the same concept. Skills may appear in
    a dedicated skills section, or be mentioned within an experience or
    internship description — use your understanding to identify them.
    Return only valid JSON matching this schema: {resume_schema}
    Do not invent information. Return only valid JSON matching this schema.

    If a string field is missing, return "".

    If a list field is missing, return [].

    Do not return null.

    Do not invent information.
    """

    resume_user_prompt = f"Parse the following resume: {resume_text}"

    message_system_resume = {"role": "system", "content": resume_system_prompt}
    message_user_resume = {"role": "user", "content": resume_user_prompt}

    messages_resume = [message_system_resume, message_user_resume]

    response_resume = client.chat.completions.create(
        model=model,
        messages=messages_resume,
        response_format=response_format
    )

    raw_resume_json = response_resume.choices[0].message.content
    resume_data = json.loads(raw_resume_json)
    print(resume_data)
    resume = Resume(**resume_data)

    return resume

# ============================================================
# PART 5 — Scoring
# ============================================================

class MatchResult(BaseModel):
    score: float
    details: dict

match_schema = MatchResult.model_json_schema()

def final_score(job, parsed_resume):
    score_prompt = f"""
    You are an HR recruiter. Compare the candidate's resume with the
    job description.

    Job Description: {job}
    Candidate Resume: {parsed_resume}

    Give me: the candidate's name, matching skills, missing skills,
    whether the experience requirement is met, an overall match
    percentage from 0 to 100, and a short final verdict explaining the
    score. Keep the response short and to the point. Return only valid
    JSON matching this schema: {match_schema}
    """

    message_score = {"role": "user", "content": score_prompt}
    messages_score = [message_score]

    response_score = client.chat.completions.create(
        model=model,
        messages=messages_score,
        response_format=response_format
    )

    raw_match_json = response_score.choices[0].message.content
    match_result_data = json.loads(raw_match_json)
    result = MatchResult(**match_result_data)

    return result

# ============================================================
# MAIN PIPELINE — Iterate Over Résumés, Score, Rank
# ============================================================

resume_folder = Path("resumes")

all_results = []

for file_path in resume_folder.iterdir():
    suffix = str(file_path).lower()
    if not (suffix.endswith(".pdf") or suffix.endswith(".docx")):
        continue

    print(f"Processing {file_path}")

    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)

    time.sleep(5)

    result = final_score(job, parsed_resume)

    time.sleep(5)

    print(result.score)

    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })

all_results.sort(key=lambda r: r["score"], reverse=True)

print("Top 2 candidates:")
for candidate in all_results[:2]:
    print(candidate)

print("Bottom 2 candidates:")
for candidate in all_results[-2:]:
    print(candidate)